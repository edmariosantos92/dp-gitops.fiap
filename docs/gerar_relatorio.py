# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2)

def set_font(run, name='Arial', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, size=14, bold=True, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def body(doc, text, size=11, space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def code_block(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, name='Courier New', size=size)
    run.font.color.rgb = RGBColor(30, 30, 30)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

AZUL = (0, 70, 127)

# ── CAPA ──────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('ATIVIDADE PRÁTICA DE DEPENDÊNCIA')
set_font(r, size=16, bold=True, color=AZUL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Infrastructure as Code (GitOps)')
set_font(r, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(60)
r = p.add_run('Pipeline GitOps para provisionamento, validação e entrega de infraestrutura imutável')
set_font(r, size=11, italic=True, color=(80, 80, 80))

for label, value in [
    ('Aluno:', 'Edmário Santos'),
    ('Disciplina:', 'Infrastructure as Code (GitOps)'),
    ('Turma:', '1TCNPZ'),
    ('Data:', '19/06/2026'),
    ('Vídeo:', 'https://diascarneiro-my.sharepoint.com/:f:/g/personal/esa_diascarneiro_com_br/IgBUmfHRiNnfRY5v3KeHWZe2AaOo-7deahO0BH5daWNzrjI?e=IoLelf'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f'{label} ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(value)
    set_font(r2, size=12)

doc.add_page_break()

# ── 1. RESUMO EXECUTIVO ───────────────────────────────────
heading(doc, '1. Resumo Executivo', size=14, color=AZUL)
body(doc,
    'Este trabalho implementou uma esteira GitOps completa sobre infraestrutura AWS, '
    'cobrindo provisionamento, entrega contínua e reconciliação declarativa. '
    'A solução utiliza Terraform para provisionar VPC, cluster EKS e repositório ECR '
    'em dois ambientes independentes (dev e prd), controlados pelas branches develop e main. '
    'Um pipeline de CI/CD no GitHub Actions valida, testa e publica automaticamente a '
    'aplicação Node.js containerizada, atualizando os manifestos Kubernetes via Kustomize. '
    'O ArgoCD monitora o repositório Git e sincroniza o estado declarado para o cluster EKS '
    'de forma automática, garantindo reconciliação contínua e detecção de drift. '
    'Foram implementadas branch protection rules que bloqueiam merges em caso de falha nos '
    'checks de CI, demonstrando governança do fluxo de mudança. '
    'Experimentos de rollback via interface do ArgoCD e de destruição de infraestrutura via '
    'pipeline automatizado completam a cobertura dos requisitos da atividade.'
)

doc.add_page_break()

# ── 2. ARQUITETURA E DECISÕES TÉCNICAS ───────────────────
heading(doc, '2. Arquitetura e Decisões Técnicas', size=14, color=AZUL)

heading(doc, '2.1 Visão Geral do Fluxo GitOps', size=12, color=AZUL, space_before=6)
body(doc, 'O fluxo de entrega segue o modelo GitOps onde o repositório Git é a fonte única de verdade:')
for step in [
    'Desenvolvedor faz push ou abre Pull Request → dispara pipeline CI/CD no GitHub Actions',
    'CI: lint, testes unitários, terraform validate/plan, tfsec (security scan)',
    'CD (apenas em push para develop ou main): build da imagem Docker, push para ECR com tag = SHA do commit',
    'Atualização automática do kustomization.yaml com a nova tag via commit do bot [skip ci]',
    'ArgoCD detecta alteração no repositório e sincroniza o manifesto para o cluster EKS',
    'Aplicação é atualizada no ambiente correspondente (dev ou prd) sem intervenção manual',
]:
    bullet(doc, step)

heading(doc, '2.2 Ferramentas Utilizadas e Justificativas', size=12, color=AZUL, space_before=8)
tools = [
    ('Terraform ~1.5',
     'Provisionamento declarativo de infraestrutura AWS (VPC, EKS, ECR). Escolhido pela maturidade, '
     'suporte nativo à AWS e backend remoto com S3 + DynamoDB para controle de estado e lock.'),
    ('GitHub Actions',
     'Pipeline CI/CD nativo do repositório, sem necessidade de servidor externo. '
     'Integração direta com eventos de PR e push, path filters e branch protection.'),
    ('Amazon EKS',
     'Kubernetes gerenciado na AWS. Elimina a gestão do control plane e integra nativamente com IAM, ECR e VPC.'),
    ('Amazon ECR',
     'Registry de imagens Docker privado na AWS. Integração nativa com EKS via IAM, '
     'sem necessidade de credenciais adicionais nos pods.'),
    ('ArgoCD',
     'Operador GitOps para Kubernetes. Monitora o repositório Git e garante que o estado do cluster '
     'reflita o estado declarado, com reconciliação automática (selfHeal) e histórico de deployments para rollback.'),
    ('Kustomize',
     'Gerenciamento de overlays por ambiente (dev/prd) sem duplicação de manifestos base. '
     'Permite customizar imagens, réplicas e variáveis por ambiente.'),
    ('tfsec',
     'Análise estática de segurança nos arquivos Terraform. '
     'Integrado ao pipeline de CI para detectar configurações inseguras antes do apply.'),
]
for tool, justificativa in tools:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f'{tool}: ')
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(justificativa)
    set_font(r2, size=11)

heading(doc, '2.3 Ambientes', size=12, color=AZUL, space_before=8)
body(doc, 'Foram criados dois ambientes lógicos com separação de estado e configuração:')
for env in [
    'dev: branch develop → cluster EKS eks-rm565486-dev → overlay k8s/overlays/dev/',
    'prd: branch main → cluster EKS eks-rm565486-prd → overlay k8s/overlays/prd/',
]:
    bullet(doc, env)

doc.add_page_break()

# ── 3. PASSO A PASSO ─────────────────────────────────────
heading(doc, '3. Passo a Passo de Implementação', size=14, color=AZUL)

heading(doc, '3.1 Estrutura do Repositório', size=12, color=AZUL, space_before=6)
for linha in [
    'dp-gitops.fiap/',
    '   .github/workflows/',
    '      ci-cd-app.yaml       # Pipeline CI/CD da aplicação',
    '      ci-cd-iac.yaml       # Pipeline CI/CD de infraestrutura',
    '      terraform-destroy.yaml  # Workflow de destruição manual',
    '   application/            # Código-fonte Node.js + Dockerfile',
    '   k8s/',
    '      base/                # Manifestos Kubernetes base',
    '      overlays/dev/        # Overlay ambiente dev',
    '      overlays/prd/        # Overlay ambiente prd',
    '      argocd/              # Definição das apps ArgoCD',
    '   terraform/              # Infraestrutura como código',
    '   docs/                   # Documentação e diagrama',
]:
    code_block(doc, linha)

heading(doc, '3.2 Provisionamento da Infraestrutura', size=12, color=AZUL, space_before=8)
body(doc,
    'A infraestrutura é provisionada automaticamente via pipeline ao fazer push para develop (dev) '
    'ou main (prd). O workflow ci-cd-iac.yaml executa Terraform Apply com o backend de estado correto '
    'para cada ambiente:'
)
for linha in [
    'terraform init -backend-config=backend-configs/dev.hcl',
    'terraform apply -auto-approve -var="env=dev"',
]:
    code_block(doc, linha)
body(doc, 'Recursos criados por ambiente:', space_before=6)
for r in [
    'VPC com subnets públicas e privadas em múltiplas AZs',
    'Cluster EKS com node group gerenciado (instância t3.medium)',
    'Repositório ECR para armazenar imagens Docker',
    'ArgoCD instalado no cluster via Terraform (Helm provider)',
]:
    bullet(doc, r)

heading(doc, '3.3 Pipeline CI/CD da Aplicação (ci-cd-app.yaml)', size=12, color=AZUL, space_before=8)
body(doc,
    'Disparado por push ou PR nas branches main/develop com alterações em application/** '
    'ou no próprio arquivo de workflow. Separado em dois jobs:'
)
body(doc, 'Job "Lint & Test" — executado em PR e push:', space_before=4)
for linha in ['npm install', 'npm run lint   # ESLint', 'npm test       # Node.js built-in test runner']:
    code_block(doc, linha)
body(doc, 'Job "Build, Push & Deploy" — executado apenas em push:', space_before=6)
for linha in [
    'docker build --build-arg IMAGE_TAG=$SHA -t $ECR_REPO:$SHA ./application',
    'docker push $ECR_REPO:$SHA',
    'kustomize edit set image rm565486-app=$ECR_REPO:$SHA',
    'git commit -m "chore(env): update image to $SHA [skip ci]"',
    'git push origin $BRANCH',
]:
    code_block(doc, linha)

heading(doc, '3.4 Pipeline CI/CD de Infraestrutura (ci-cd-iac.yaml)', size=12, color=AZUL, space_before=8)
body(doc, 'Disparado por push ou PR em terraform/** ou no próprio workflow.')
body(doc, 'Job CI (apenas em PR — evita conflito de state lock em push):')
for c in [
    'terraform fmt -check -recursive',
    'terraform init -backend-config=backend-configs/$ENV.hcl',
    'terraform validate',
    'terraform plan -var="env=$ENV" -no-color',
    'tfsec (security scan, soft_fail: true)',
]:
    bullet(doc, c)
body(doc, 'Job CD (apenas em push para develop ou main):', space_before=4)
bullet(doc, 'terraform apply -auto-approve -var="env=dev|prd"')

heading(doc, '3.5 Configuração do ArgoCD', size=12, color=AZUL, space_before=8)
body(doc,
    'O ArgoCD foi instalado via Terraform (Helm provider) e configurado com '
    'sincronização automática e self-heal. Cada ambiente possui uma Application separada:'
)
for linha in [
    'apiVersion: argoproj.io/v1alpha1',
    'kind: Application',
    'spec:',
    '  source:',
    '    repoURL: https://github.com/edmariosantos92/dp-gitops.fiap',
    '    targetRevision: main  # ou develop para dev',
    '    path: k8s/overlays/prd',
    '  syncPolicy:',
    '    automated:',
    '      prune: true      # Remove recursos deletados do Git',
    '      selfHeal: true   # Reverte alterações manuais no cluster',
]:
    code_block(doc, linha)

heading(doc, '3.6 Branch Protection Rules', size=12, color=AZUL, space_before=8)
body(doc,
    'Configuradas via GitHub API nas branches develop e main para exigir que o check '
    '"Lint & Test" passe antes de qualquer merge. Bloqueia force push e exclusão das branches protegidas.'
)

doc.add_page_break()

# ── 4. TESTES EXECUTADOS ──────────────────────────────────
heading(doc, '4. Testes Executados', size=14, color=AZUL)

heading(doc, '4.1 Pipeline CI/CD com sucesso — PR #10 (feat/version-endpoint → develop)', size=12, color=AZUL, space_before=6)
body(doc, 'Resultado esperado: jobs Lint & Test e Build Push & Deploy com status success.')
body(doc,
    'Resultado obtido: Lint & Test concluído em 14s. Build, Push & Deploy concluído em 29s. '
    'Imagem publicada no ECR com tag = SHA do commit. kustomization.yaml atualizado '
    'automaticamente e ArgoCD sincronizado com o novo deployment.',
    space_before=4
)

heading(doc, '4.2 PR bloqueado por falha de CI — PR #9 (test/ci-block-evidence → develop)', size=12, color=AZUL, space_before=8)
body(doc, 'Resultado esperado: merge bloqueado pelo GitHub quando o check "Lint & Test" falha.')
body(doc,
    'Resultado obtido: CI falhou com AssertionError: "ok" !== "healthy" no teste GET /health. '
    'Branch protection rule impediu o merge com mensagem "Required status check Lint & Test is failing". '
    'Botão "Merge pull request" desabilitado.',
    space_before=4
)
body(doc,
    'Análise: o bloqueio só passou a funcionar após configuração explícita das branch protection rules. '
    'O GitHub não bloqueia merges automaticamente apenas por falha de CI — é necessária a configuração '
    'de required status checks na proteção da branch.',
    space_before=4
)

heading(doc, '4.3 Drift e Reconciliação Automática (ArgoCD selfHeal)', size=12, color=AZUL, space_before=8)
body(doc, 'Resultado esperado: ArgoCD detecta alteração manual no cluster e reverte para o estado declarado no Git.')
body(doc,
    'Resultado obtido: Após deleção manual de recurso via UI do ArgoCD, o app entrou em estado '
    'OutOfSync e foi automaticamente reconciliado para Synced em menos de 3 minutos, '
    'sem intervenção manual.',
    space_before=4
)

heading(doc, '4.4 Rollback via ArgoCD UI', size=12, color=AZUL, space_before=8)
body(doc, 'Resultado esperado: ArgoCD permite reverter para versão anterior via histórico de syncs.')
body(doc,
    'Resultado obtido: Rollback executado via "History and Rollback" na UI do ArgoCD. '
    'App implantou versão anterior com sucesso. Com selfHeal ativo, o ArgoCD reconvergiu '
    'para o HEAD do repositório automaticamente após o rollback manual, demonstrando '
    'tanto o rollback quanto a reconciliação declarativa.',
    space_before=4
)

heading(doc, '4.5 Destruição de Infraestrutura (terraform-destroy.yaml)', size=12, color=AZUL, space_before=8)
body(doc, 'Resultado esperado: workflow destroy executa terraform destroy nos ambientes dev e prd sem erros.')
body(doc,
    'Primeiro run falhou com erro "Error acquiring the state lock" — state lock travado por pipeline '
    'cancelado anteriormente (OperationTypePlan pendente no DynamoDB). '
    'Corrigido com adição de step para deletar o item de lock no DynamoDB antes do destroy. '
    'Segunda execução disparada com sucesso para ambos os ambientes em paralelo.',
    space_before=4
)

doc.add_page_break()

# ── 5. INVENTÁRIO DE RECURSOS ─────────────────────────────
heading(doc, '5. Inventário de Recursos e Instruções de Limpeza', size=14, color=AZUL)

heading(doc, '5.1 Recursos AWS (por ambiente dev e prd)', size=12, color=AZUL, space_before=6)
recursos = [
    ('VPC', 'rm565486-vpc-dev / rm565486-vpc-prd'),
    ('Subnets', '3 públicas + 3 privadas por VPC'),
    ('Internet Gateway', '1 por VPC'),
    ('NAT Gateway', '1 por VPC'),
    ('EKS Cluster', 'eks-rm565486-dev / eks-rm565486-prd (Kubernetes 1.33)'),
    ('EKS Node Group', 't3.medium, mín: 1 / máx: 2 nós'),
    ('ECR Repository', 'rm565486-app (compartilhado)'),
    ('S3 Bucket', 'rm565486-tfstate (estado Terraform)'),
    ('DynamoDB Table', 'rm565486-tflock (state locking)'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Recurso', 'Nome / Configuração']):
    p = hdr[i].paragraphs[0]
    run = p.add_run(h)
    set_font(run, size=10, bold=True)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'D6E4F0')
    hdr[i]._tc.get_or_add_tcPr().append(shading)
for recurso, nome in recursos:
    row = table.add_row().cells
    for i, val in enumerate([recurso, nome]):
        p = row[i].paragraphs[0]
        run = p.add_run(val)
        set_font(run, size=10)

heading(doc, '5.2 Recursos Kubernetes', size=12, color=AZUL, space_before=10)
for r in [
    'Namespace: rm565486-dev e rm565486-prd',
    'Deployment: rm565486-app (1 réplica por ambiente)',
    'Service: rm565486-app (ClusterIP)',
    'Namespace argocd com todos os componentes do ArgoCD',
    'ArgoCD Application: rm565486-app-dev e rm565486-app-prd',
]:
    bullet(doc, r)

heading(doc, '5.3 Instruções de Destruição', size=12, color=AZUL, space_before=8)
body(doc, 'A destruição é realizada via workflow manual no GitHub Actions:')
for passo in [
    'Acessar o repositório no GitHub → aba Actions',
    'Selecionar o workflow "Terraform Destroy"',
    'Clicar em "Run workflow" e selecionar o ambiente (dev ou prd)',
    'Confirmar execução — o workflow exclui o state lock pendente e executa terraform destroy -auto-approve',
    'Repetir para o outro ambiente',
]:
    bullet(doc, passo)

doc.add_page_break()

# ── 6. CONCLUSÃO TÉCNICA ──────────────────────────────────
heading(doc, '6. Conclusão Técnica', size=14, color=AZUL)

heading(doc, '6.1 Resultados Alcançados', size=12, color=AZUL, space_before=6)
for r in [
    'Pipeline GitOps completo implementado e validado nos dois ambientes (dev e prd)',
    'Separação de responsabilidades entre pipeline de aplicação e de infraestrutura',
    'Reconciliação automática via ArgoCD com selfHeal e prune ativos',
    'Governança de mudanças com branch protection bloqueando merges com falha de CI',
    'Rastreabilidade total: versão da imagem em produção reflete o SHA do commit no Git',
    'Experimentos de drift, rollback e destruição documentados com evidências de execução',
]:
    bullet(doc, r)

heading(doc, '6.2 Limitações', size=12, color=AZUL, space_before=8)
for r in [
    'Cluster EKS sem Ingress Controller configurado — aplicação não possui acesso HTTP externo',
    'tfsec configurado com soft_fail: true, não bloqueando o pipeline em violações de segurança',
    'Não há ambiente de homologação intermediário entre dev e prd',
    'ArgoCD sem autenticação SSO — apenas usuário admin local',
    'Credenciais AWS armazenadas como secrets estáticos no GitHub',
]:
    bullet(doc, r)

heading(doc, '6.3 Riscos', size=12, color=AZUL, space_before=8)
for r in [
    'State lock travado em caso de pipeline cancelado — mitigado com step de force-unlock no workflow de destroy',
    'Commit [skip ci] no HEAD da branch pode suprimir triggers de PR — identificado e documentado durante a atividade',
    'Segredos AWS expostos em caso de comprometimento do repositório GitHub',
]:
    bullet(doc, r)

heading(doc, '6.4 Melhorias Futuras', size=12, color=AZUL, space_before=8)
for r in [
    'Implementar IRSA (IAM Roles for Service Accounts) para eliminar credenciais estáticas',
    'Adicionar Ingress Controller (AWS ALB) com certificado TLS para exposição da aplicação',
    'Configurar tfsec como bloqueante (remover soft_fail)',
    'Implementar ambiente de staging entre dev e prd',
    'Adicionar monitoramento com Prometheus e Grafana no cluster EKS',
    'Substituir workflow_dispatch do destroy por environment protection rules com aprovação obrigatória',
]:
    bullet(doc, r)

doc.add_page_break()

# ── 7. REFERÊNCIAS E DECLARAÇÃO ───────────────────────────
heading(doc, '7. Referências e Declaração de Autoria', size=14, color=AZUL)

heading(doc, '7.1 Referências', size=12, color=AZUL, space_before=6)
refs = [
    'HashiCorp. Terraform Documentation. https://developer.hashicorp.com/terraform',
    'Amazon Web Services. Amazon EKS User Guide. https://docs.aws.amazon.com/eks',
    'Argo CD. Declarative GitOps CD for Kubernetes. https://argo-cd.readthedocs.io',
    'GitHub. GitHub Actions Documentation. https://docs.github.com/actions',
    'Kubernetes. Kustomize Reference. https://kubectl.docs.kubernetes.io/references/kustomize',
    'Aqua Security. tfsec - Security Scanner for Terraform. https://github.com/aquasecurity/tfsec',
    'Weaveworks. GitOps Principles. https://www.gitops.tech',
    'Repositório da atividade. https://github.com/edmariosantos92/dp-gitops.fiap',
    'Vídeo de demonstração. https://diascarneiro-my.sharepoint.com/:f:/g/personal/esa_diascarneiro_com_br/IgBUmfHRiNnfRY5v3KeHWZe2AaOo-7deahO0BH5daWNzrjI?e=IoLelf',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'[{i}] {ref}')
    set_font(r, size=10)

heading(doc, '7.2 Declaração de Autoria', size=12, color=AZUL, space_before=10)
body(doc,
    'Declaro que este trabalho foi desenvolvido integralmente por mim, Edmário Santos, '
    'aluno da turma 1TCNPZ, como parte da atividade prática de dependência da disciplina '
    'Infrastructure as Code (GitOps). Todo o código, configurações e documentação '
    'presentes neste relatório e no repositório GitHub são de minha autoria, com uso de '
    'ferramentas de inteligência artificial como suporte à implementação e documentação.'
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Edmário Santos  |  1TCNPZ  |  19/06/2026')
set_font(r, size=11)

doc.save('C:/Users/edmar/Desktop/DP/docs/relatorio.docx')
print('OK: docs/relatorio.docx gerado com sucesso')
